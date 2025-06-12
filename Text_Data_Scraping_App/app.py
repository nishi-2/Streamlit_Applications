import streamlit as st
import requests
from bs4 import BeautifulSoup
from scrape import get_data
from docx import Document  # Required to create Word document
from io import BytesIO     # Required to create an in-memory file

st.title("Web Scraping")
st.write("Enter a URL to scrape its content - ")

url_input = st.text_input("Enter URL")

if st.button("Scrape"):
    if url_input:
        content = get_data(url_input)
        if content:
            st.subheader("Scraped Content:")
            st.write(content)

            # Create Word Document in memory
            doc = Document()
            doc.add_heading('Scraped Web Content', 0)
            doc.add_paragraph(content)

            # Save to in-memory buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Add download button
            st.download_button(
                label="Download as Word File",
                data=buffer,
                file_name="scraped_content.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("No content found or an error occurred.")
    else:
        st.warning("Please enter a valid URL.")
